import logging
import re
import tempfile
from io import BytesIO
from pathlib import Path
from typing import List

from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .latex_compiler import compile_latex

logger = logging.getLogger(__name__)


class PDFService:
    @staticmethod
    def _remove_glyph_to_unicode_lines(latex_text: str) -> str:
        cleaned = re.sub(r'^\s*\\input\{glyphtounicode\}\s*$', '', latex_text, flags=re.MULTILINE)
        cleaned = re.sub(r'^\s*\\pdfgentounicode\s*=\s*1\s*$', '', cleaned, flags=re.MULTILINE)
        cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
        return cleaned

    @staticmethod
    def _remove_fontawesome_dependency(latex_text: str) -> str:
        cleaned = re.sub(
            r'^\s*\\usepackage(?:\[[^\]]*\])?\{fontawesome5\}\s*$',
            '',
            latex_text,
            flags=re.MULTILINE,
        )
        # Replace common icon commands with plain spacing so contact text remains readable.
        cleaned = re.sub(r'\\fa[A-Za-z]+\s*\\?', ' ', cleaned)
        cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
        return cleaned

    @staticmethod
    def compile_latex_to_pdf(latex_text, timeout_seconds=180):
        """Compile LaTeX source into a PDF (Overleaf-like build)."""
        if not latex_text or not str(latex_text).strip():
            raise ValueError("LaTeX content is empty")

        source_text = str(latex_text)
        candidates: list[tuple[str, str]] = [('original', source_text)]
        if r'\input{glyphtounicode}' in source_text or r'\pdfgentounicode=1' in source_text:
            candidates.append(('without_glyph_to_unicode', PDFService._remove_glyph_to_unicode_lines(source_text)))
        if r'\usepackage{fontawesome5}' in source_text or r'\fa' in source_text:
            candidates.append(('without_fontawesome', PDFService._remove_fontawesome_dependency(source_text)))
        if len(candidates) > 1:
            candidates.append((
                'without_glyph_to_unicode_and_fontawesome',
                PDFService._remove_fontawesome_dependency(PDFService._remove_glyph_to_unicode_lines(source_text)),
            ))

        last_error: Exception | None = None
        try:
            for variant_name, candidate_text in candidates:
                with tempfile.TemporaryDirectory() as temp_dir:
                    temp_path = Path(temp_dir)
                    tex_path = temp_path / 'resume.tex'
                    tex_path.write_text(candidate_text, encoding='utf-8')
                    try:
                        pdf_path = Path(
                            compile_latex(
                                tex_path=str(tex_path),
                                output_dir=str(temp_path),
                                timeout_seconds=timeout_seconds,
                            )
                        )
                    except Exception as exc:
                        logger.warning("LaTeX compile attempt failed (%s): %s", variant_name, str(exc))
                        last_error = exc
                        continue

                    if variant_name != 'original':
                        logger.warning("LaTeX compile succeeded after fallback variant: %s", variant_name)

                    pdf_buffer = BytesIO(pdf_path.read_bytes())
                    pdf_buffer.seek(0)
                    return pdf_buffer

            if last_error:
                raise last_error
            raise RuntimeError("LaTeX compilation failed for all variants.")
        except Exception as e:
            logger.error(f"Error compiling LaTeX to PDF: {str(e)}")
            raise

    @staticmethod
    def extract_text_from_pdf(file):
        """Extract text from PDF file."""
        try:
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            if not text.strip():
                raise ValueError("PDF appears to be empty or scanned (no text layer)")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise ValueError(f"Failed to read PDF file: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file):
        """Extract text from DOCX file."""
        try:
            doc = Document(file)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            if not text.strip():
                raise ValueError("DOCX file appears to be empty")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise ValueError(f"Failed to read DOCX file: {str(e)}")

    @staticmethod
    def extract_text(file):
        """Extract text from file based on extension."""
        if not file:
            raise ValueError("No file provided")

        filename = file.name.lower()

        try:
            if filename.endswith('.pdf'):
                return PDFService.extract_text_from_pdf(file)
            if filename.endswith('.docx'):
                return PDFService.extract_text_from_docx(file)
            if filename.endswith('.txt'):
                content = file.read()
                text = content.decode('utf-8') if isinstance(content, bytes) else content
                if not text.strip():
                    raise ValueError("Text file is empty")
                return text
            if filename.endswith('.tex'):
                content = file.read()
                if isinstance(content, bytes):
                    try:
                        text = content.decode('utf-8')
                    except UnicodeDecodeError:
                        text = content.decode('latin-1')
                else:
                    text = content
                if not text.strip():
                    raise ValueError("TeX file is empty")
                return text
            raise ValueError("Unsupported file format. Please upload PDF, DOCX, TXT, or TEX")
        except UnicodeDecodeError:
            raise ValueError("File encoding error. Please ensure the file is valid")
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            raise

    @staticmethod
    def generate_resume_pdf(resume_data):
        """Generate PDF from structured resume data."""
        if not isinstance(resume_data, dict):
            raise ValueError("Invalid resume data format")

        try:
            pdf_file = BytesIO()
            doc = SimpleDocTemplate(pdf_file, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            name = resume_data.get('name', 'Your Name')
            story.append(Paragraph(name, styles['Title']))

            email = resume_data.get('email', '')
            phone = resume_data.get('phone', '')
            contact = f"{email} | {phone}" if email and phone else email or phone
            if contact:
                story.append(Paragraph(contact, styles['Normal']))
            story.append(Spacer(1, 12))

            skills = resume_data.get('skills', [])
            if skills and isinstance(skills, list):
                story.append(Paragraph('Skills', styles['Heading2']))
                story.append(Paragraph(', '.join(str(s) for s in skills), styles['Normal']))
                story.append(Spacer(1, 12))

            experience = resume_data.get('experience', [])
            if experience and isinstance(experience, list):
                story.append(Paragraph('Experience', styles['Heading2']))
                for exp in experience:
                    if isinstance(exp, dict):
                        title = exp.get('title', '')
                        company = exp.get('company', '')
                        duration = exp.get('duration', '')
                        story.append(Paragraph(f"{title} - {company} ({duration})", styles['Heading3']))

                        responsibilities = exp.get('responsibilities', [])
                        if isinstance(responsibilities, list):
                            for resp in responsibilities:
                                story.append(Paragraph(f"- {resp}", styles['Normal']))
                        story.append(Spacer(1, 6))

            education = resume_data.get('education', [])
            if education and isinstance(education, list):
                story.append(Paragraph('Education', styles['Heading2']))
                for edu in education:
                    if isinstance(edu, dict):
                        degree = edu.get('degree', '')
                        institution = edu.get('institution', '')
                        year = edu.get('year', '')
                        story.append(Paragraph(f"{degree} - {institution} ({year})", styles['Normal']))

            doc.build(story)
            pdf_file.seek(0)
            return pdf_file
        except Exception as e:
            logger.error(f"Error generating resume PDF: {str(e)}")
            raise ValueError(f"Failed to generate PDF: {str(e)}")

    @staticmethod
    def generate_cover_letter_pdf(content, name):
        """Generate cover letter PDF."""
        if not content or not content.strip():
            raise ValueError("Cover letter content is empty")

        try:
            pdf_file = BytesIO()
            doc = SimpleDocTemplate(pdf_file, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            display_name = (name or '').strip()
            if display_name:
                story.append(Paragraph(display_name, styles['Heading1']))
                story.append(Spacer(1, 12))

            for para in content.split('\n'):
                if para.strip():
                    story.append(Paragraph(para, styles['Normal']))
                    story.append(Spacer(1, 6))

            doc.build(story)
            pdf_file.seek(0)
            return pdf_file
        except Exception as e:
            logger.error(f"Error generating cover letter PDF: {str(e)}")
            raise ValueError(f"Failed to generate cover letter PDF: {str(e)}")

    @staticmethod
    def generate_text_pdf(title, content):
        """Generate generic text-based PDF."""
        if not content or not content.strip():
            raise ValueError("PDF content is empty")

        try:
            pdf_file = BytesIO()
            doc = SimpleDocTemplate(pdf_file, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            title_text = (title or '').strip()
            if title_text:
                story.extend([Paragraph(title_text, styles['Title']), Spacer(1, 16)])

            for raw_line in content.split('\n'):
                line = raw_line.strip()
                if line:
                    story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 8))

            doc.build(story)
            pdf_file.seek(0)
            return pdf_file
        except Exception as e:
            logger.error(f"Error generating text PDF: {str(e)}")
            raise ValueError(f"Failed to generate PDF: {str(e)}")

    @staticmethod
    def _escape_latex_text(value: str) -> str:
        escaped = str(value or '')
        replacements = {
            '\\': r'\textbackslash{}',
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\textasciicircum{}',
        }
        for source, target in replacements.items():
            escaped = escaped.replace(source, target)
        return escaped

    @staticmethod
    def _to_paragraphs(content: str) -> List[str]:
        normalized = (content or '').replace('\r\n', '\n').replace('\r', '\n')
        return [part.strip() for part in normalized.split('\n\n') if part.strip()]

    @staticmethod
    def build_cover_letter_latex(content, name=''):
        if not content or not content.strip():
            raise ValueError("Cover letter content is empty")

        header = PDFService._escape_latex_text((name or '').strip())
        paragraphs = PDFService._to_paragraphs(content)
        if not paragraphs:
            raise ValueError("Cover letter content is empty")

        body_blocks = []
        for paragraph in paragraphs:
            raw_lines = [line.strip() for line in paragraph.split('\n') if line.strip()]
            escaped_lines = [PDFService._escape_latex_text(line) for line in raw_lines]
            if not escaped_lines:
                continue
            block = " \\\\\n".join(escaped_lines)
            body_blocks.append(f"{block}\n\\par")
        body = "\n\n".join(body_blocks)
        header_block = f"\\textbf{{{header}}}\n\n" if header else ""
        return rf"""\documentclass[11pt]{{letter}}
\usepackage[margin=1in]{{geometry}}
\usepackage[T1]{{fontenc}}
\usepackage[utf8]{{inputenc}}
\usepackage{{lmodern}}
\setlength{{\parindent}}{{0pt}}
\setlength{{\parskip}}{{10pt}}
\begin{{document}}
\raggedright
{header_block}
{body}
\end{{document}}
"""

    @staticmethod
    def generate_cover_letter_pdf_via_latex(content, name=''):
        latex_text = PDFService.build_cover_letter_latex(content=content, name=name)
        return PDFService.compile_latex_to_pdf(latex_text)

    @staticmethod
    def generate_cover_letter_docx(content, name=''):
        if not content or not content.strip():
            raise ValueError("Cover letter content is empty")

        try:
            document = Document()
            display_name = (name or '').strip()
            if display_name:
                document.add_paragraph(display_name)

            for paragraph in PDFService._to_paragraphs(content):
                document.add_paragraph(paragraph)

            docx_file = BytesIO()
            document.save(docx_file)
            docx_file.seek(0)
            return docx_file
        except Exception as e:
            logger.error(f"Error generating cover letter DOCX: {str(e)}")
            raise ValueError(f"Failed to generate cover letter DOCX: {str(e)}")
